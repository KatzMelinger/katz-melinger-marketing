"""Extract text from the Katz Melinger SOPs into TSV files for analysis."""
from pathlib import Path
from docx import Document

SRC = Path(r"C:\Users\KM\Desktop\Software\katz-melinger-marketing\docs\sales-training\Katz Melinger SOPS")
OUT = Path(__file__).parent / "sops_extract"
OUT.mkdir(exist_ok=True)

for fp in sorted(SRC.glob("*.docx")):
    doc = Document(fp)
    lines = []
    for p in doc.paragraphs:
        t = p.text.strip()
        if t:
            lines.append(t)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                lines.append(" | ".join(cells))
    out = OUT / (fp.stem + ".txt")
    out.write_text("\n".join(lines), encoding="utf-8")
    print(f"{fp.name}: {len(lines)} lines, {out.stat().st_size:,} bytes")
